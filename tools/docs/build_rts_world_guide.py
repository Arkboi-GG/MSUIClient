from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "artifacts" / "RTS_World_Architecture_and_Game_Loop.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"
MUTED = "666666"


def set_run(run, size=11, bold=False, italic=False, color="000000"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "left" if side == "start" else "right" if side == "end" else side
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_table_geometry(table, widths):
    assert sum(widths) == 9360
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def style_cell_text(cell, size=9.5, bold=False, color="000000", align=WD_ALIGN_PARAGRAPH.LEFT):
    for p in cell.paragraphs:
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            set_run(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
        shade(table.rows[0].cells[idx], LIGHT_BLUE)
        style_cell_text(table.rows[0].cells[idx], size=9.2, bold=True, color=INK)
    set_repeat_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cells[idx].text = str(text)
            style_cell_text(cells[idx], size=font_size)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    return table


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True, color=INK)
        set_run(p.add_run(text[len(bold_lead):]))
    else:
        set_run(p.add_run(text))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_run(p.add_run(item))


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_run(p.add_run(item))


def add_callout(doc, label, text, tone="blue"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT if tone == "blue" else "FFF7E6" if tone == "gold" else "FCEBEC")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), DARK_BLUE if tone == "blue" else GOLD if tone == "gold" else RED)
    borders.append(left)
    p_pr.append(borders)
    set_run(p.add_run(label + "  "), bold=True,
            color=DARK_BLUE if tone == "blue" else GOLD if tone == "gold" else RED)
    set_run(p.add_run(text), size=10.5)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.add_run("RTS WORLD ARCHITECTURE  |  TECHNICAL FIELD GUIDE"),
            size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.add_run("MSUIClient / MangosSuperUI / SuperUI-Core  |  2026-08-16  |  "),
            size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_cover(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(90)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_run(kicker.add_run("TIER-2 RTS MATCH LAYER"), size=11, bold=True, color=GOLD)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run(title.add_run("RTS World Architecture"), size=28, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    set_run(subtitle.add_run("How the world state, Core, web creator, client, and effective game loop connect"),
            size=14, color=DARK_BLUE)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(70)
    set_run(meta.add_run("R1-R4 system model  |  R3 implementation handoff  |  16 August 2026"),
            size=10, italic=True, color=MUTED)
    add_callout(doc, "SYSTEM TRUTH",
        "R3 is not end-to-end complete until the web creator, authoritative Linux Core, authored pilot assets, and Nico-operated live acceptance all pass. This work completes the MSUIClient R3 contract and presentation lane; it does not deploy or mutate live state.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Fight -> Honor -> Heroes. Hold -> Capacity -> Scale. Clear -> Buff. Push -> Victory."),
            size=12, bold=True, color=INK)
    doc.add_page_break()


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)
    add_cover(doc)

    doc.add_heading("1. The effective game loop", level=1)
    add_body(doc, "The RTS world is not a separate minigame layered over an idle MMO. It turns the existing vanilla world into the progression engine, battlefield, logistics map, hero roster, and objective network. Each system feeds another system, but authority remains server-side and save-bound.")
    add_steps(doc, [
        "Fight in the ordinary world. Bot-versus-bot kills add faction Honor without creating vanilla character Honor; human kills retain stock honor behavior.",
        "Spend faction Honor to declare, upgrade, and revive bot heroes. Heroes gain persistent level, scale, damage, and death state.",
        "Capture hubs to control zones. Territory immediately changes guards, standing supplies, graveyard preference, and the number of heroes the faction may declare.",
        "Use the accelerated vanilla world to level, gear, craft, earn money, and strengthen the army. Ore, Skins, and Herbs are standing strategic capacity, not banked currency.",
        "Clear enemy-accessible dungeons to claim faction-wide buffs and deny the controlling faction entry. R4 adds this objective loop after R3 is proven.",
        "Converge on capitals and faction commanders. The final victory/defeat mechanic remains an unphased design gap and must not be implied by R3 or R4.",
    ])
    add_callout(doc, "PLAYER DECISION LOOP",
        "Choose where to fight, which hub to contest, whether to invest Honor in more heroes or stronger heroes, when to defend territory, and when to convert map position into dungeon buffs and a capital push.")

    doc.add_heading("2. One world, three codebases, four authorities", level=1)
    add_table(doc,
        ["Piece", "Owns", "Code effect"],
        [
            ["MangosSuperUI", "Artifact creation and stopped-world resume", "Creates the nine-table overlay, seeds genesis state, authors catalog-owned WorldDatabase content, validates source artifacts, and publishes parked snapshots."],
            ["SuperUI-Core", "Live mechanics", "Boot-validates rules, runs capture/hero/dungeon state machines, commits runtime state, updates guards and buffs, selects graveyards, and emits authoritative wire snapshots."],
            ["MSUIClient", "Presentation and commander intent", "Strictly parses snapshots, displays the world and campaign, sends only permitted actions, and never derives authoritative control, supply, capacity, or capture completion."],
            ["CharacterDatabase", "Save-bound truth", "Owns mode, rules, Honor, heroes, zone controllers, and dungeon controllers. These rows travel with the selected RTS save."],
            ["WorldDatabase", "Shared authored content", "Owns banners, spawns, guard creatures, game events, and later buff/loot content. It is not the live match-state authority."],
            ["Nico", "Live lifecycle", "Installs, deploys, creates/loads/swaps saves, controls services, and performs live acceptance. Agents stop at source/build verification."],
        ], [1700, 2500, 5160], font_size=8.8)

    doc.add_heading("3. Save and runtime lifecycle", level=1)
    doc.add_heading("3.1 Create New RTS World", level=2)
    add_steps(doc, [
        "MangosSuperUI inspects stock Characters, World, Realmd, Core configuration, and catalog-owned ID ranges before creating output.",
        "It clones the source artifacts into a staged snapshot. No live database is modified by artifact construction.",
        "It creates exactly the nine RTS overlay tables in the cloned CharacterDatabase, writes the R3 profile and rules, seeds zeroed faction state, hero state, initial zone controllers, and empty dungeon control.",
        "It appends exact, signature-checked banner, guard, and game-event content to the cloned WorldDatabase artifact.",
        "It validates the output and publishes a parked World State. Nico alone chooses whether to load it.",
    ])
    doc.add_heading("3.2 Clean boot", level=2)
    add_steps(doc, [
        "Core loads the selected save header before bot admission and reads the ruleset once. Production rule changes are boot-time authoritative.",
        "Each module independently requests activation, validates its complete contract, and latches enabled only on success.",
        "A malformed R3 contract disables Territory only; valid Honor, Heroes, and faction-control modules continue.",
        "The first territory tick reconciles authored guard events to the committed controller. Core does not start events during world loading.",
    ])
    doc.add_heading("3.3 Resume", level=2)
    add_body(doc, "Resume performs semantic preflight before any restore, refreshes only web-managed configuration and rules, and preserves Core-owned Honor, heroes, and committed zone controllers. It performs no DDL against a running server. For the R3 pilot, Create-New-from-stock is the supported path; R2-to-R3 promotion is deferred.")

    doc.add_heading("4. The nine-table RTS overlay", level=1)
    add_table(doc,
        ["Table", "Role", "Writer"],
        [
            ["superui_worldstate", "Mode plus scalar keys: rates, module gates, cooldowns, caps", "Web while stopped; Core reads at boot"],
            ["superui_rules_zone", "Zone standing Ore, Skins, Herbs", "Web"],
            ["superui_rules_hub", "Hub, zone, banner, events, capture time, initial controller", "Web"],
            ["superui_rules_hero", "Hero levels, Honor costs, revive fees, aura spell", "Web"],
            ["superui_rules_dungeon", "Dungeon final boss, buff spell, loot item count", "Web"],
            ["superui_faction", "Persistent faction Honor pools", "Web genesis; Core runtime"],
            ["superui_heroes", "Persistent declared hero roster and death state", "Web genesis; Core runtime"],
            ["superui_zone_control", "One committed controller per configured zone", "Web genesis; Core capture commits"],
            ["superui_dungeon_control", "Persistent dungeon controllers", "Web genesis; Core in R4"],
        ], [2800, 4260, 2300], font_size=8.7)
    add_callout(doc, "AUTHORITY LAW",
        "The committed superui_zone_control row outranks memory. Supply, capacity, guards, graveyards, and wire state are projections of committed controllers; they are never independent truth.")

    doc.add_heading("5. R1 foundation: the match can exist safely", level=1)
    add_bullets(doc, [
        "Boot-time match header and module gates distinguish ordinary MMO behavior from RTS behavior.",
        "Progression and drop multipliers reuse stock VMaNGOS configuration paths.",
        "Faction bot-admission caps are applied before the bot manager admits a new bot.",
        "The RTS state packet reserves faction, hero, territory, and dungeon blocks behind explicit module flags.",
        "Every custom hook follows two gates: loaded RTS World State, then validated feature module.",
    ])
    add_callout(doc, "CURRENT STATUS", "R1 was partially owner-validated. Clean boot, scaling, character creation, and Commander state were observed; both faction admission mappings still need an explicit owner check.", tone="gold")

    doc.add_heading("6. R2: fight, Honor, heroes, and direct control", level=1)
    add_body(doc, "R2 makes combat produce faction investment choices. Honor is a faction pool. Only eligible faction AiBots may become heroes; the server remains authoritative for cost, capacity, eligibility, death, and revival.")
    add_table(doc,
        ["Event", "Authoritative effect", "Client effect"],
        [
            ["Bot-versus-bot kill", "Core awards weighted faction Honor and leaves character_honor_cp unchanged", "Commander refresh shows the new faction pool"],
            ["Declare", "Spend Honor, create persistent hero row, apply level-1 aura if below cap", "Offer Declare only as an affordance; show result and refreshed roster"],
            ["Upgrade", "Spend Honor, increment persistent hero level, replace aura", "Keep available even when the faction is over its territory cap"],
            ["Hero death", "Preserve slot and level; mark dead", "Show DEAD and offer Revive"],
            ["Revive", "Spend fee, clear dead state, resurrect using valid destination policy", "Keep available over cap; display server verdict"],
            ["Take control", "Core checks faction, presence, state, and singular control authority", "Faction roster supports selection; local UI grants no authority"],
        ], [1600, 4400, 3360], font_size=8.8)
    add_callout(doc, "CURRENT STATUS", "R2 is source/build complete across the three codebases and awaits Nico-operated deployment and live play-testing.", tone="gold")

    doc.add_heading("7. R3 territory: hold, capacity, and forward pressure", level=1)
    doc.add_heading("7.1 Pilot topology", level=2)
    add_table(doc,
        ["Hub", "Zone", "Purpose"],
        [
            ["Sentinel Hill", "Westfall", "Alliance-side pilot and Eastern Kingdoms graveyard/guard proof"],
            ["The Crossroads", "The Barrens", "Horde-side pilot and Kalimdor proof"],
            ["Tarren Mill", "Hillsbrad Foothills", "Contested-region proof and third-zone capacity behavior"],
        ], [2300, 2300, 4760], font_size=9.2)
    add_body(doc, "R3 v1 supports one capturable hub per top-level zone. The banner is a permanent neutral GOOBER; faction identity is communicated through swapped guards and client territory color, not physical flag replacement.")

    doc.add_heading("7.2 Capture state machine", level=2)
    add_steps(doc, [
        "A player uses a recognized configured banner. The map-thread hook validates both gates and enqueues only; it does not mutate shared territory state.",
        "The main-thread tick starts a neutral assault or enemy assault. A defending incumbent interaction cancels the assault; a competing neutral attacker replaces the prior attacker under the frozen rule.",
        "The server advances the authoritative timer and broadcasts packed local world-state progress at phase changes and at most once per displayed second.",
        "When capture time completes, Core snapshots the proposed controller and releases the territory lock before the database call.",
        "Core synchronously updates superui_zone_control. If the write fails, ownership, guards, supply, capacity, graveyards, and wire remain unchanged and the commit can be retried.",
        "Only after persistence succeeds does Core publish the new controller, reconcile guard events, recompute projections, and enter the configured cooldown.",
    ])
    add_callout(doc, "PERSIST BEFORE PUBLISH", "A database failure can delay a flip, but it cannot create a ghost controller visible only in memory or on the client.")

    doc.add_heading("7.3 What one controller flip changes", level=2)
    add_table(doc,
        ["Projection", "Calculation / behavior", "Affected code"],
        [
            ["Guards", "Exactly the desired faction event is active; neutral may mean neither", "SuiTerritory tick + GameEventMgr reconciliation"],
            ["Standing supply", "Sum each controlled zone's ore/skins/herbs allotment; saturate at INT32_MAX", "SuiTerritory derived snapshot -> RTS faction row -> Commander rail"],
            ["Hero capacity", "floor(controlled zones / zones_per_hero_slot); pilot default ratio 1", "SuiRts effective-cap facade -> SuiHero Declare gate -> client over-cap affordance"],
            ["Existing heroes", "Never demoted or removed when territory is lost", "Upgrade and Revive remain valid; only another Declare is blocked"],
            ["Graveyards", "Prefer nearest friendly controlled-zone graveyard on the same map", "Player repop, ordinary AiBot resurrection, paid hero revive; vanilla fallback on no candidate"],
            ["Strategic map", "Controller and contested bit remain visible even with zero population", "Zone-intel configured-zone union -> typed client snapshot -> tint/badge/pulse"],
        ], [1700, 4300, 3360], font_size=8.5)

    doc.add_heading("7.4 R3 client implementation completed in this workspace", level=2)
    add_bullets(doc, [
        "RtsWire now owns typed zone/controller/unit parsing, legacy stride-8 compatibility, future-tail skipping, duplicate/flag/finite-value validation, and atomic snapshot publication.",
        "The packed capture world state is strictly decoded against the golden 0x002A5DE9 vector and malformed states are hidden without interfering with quest macro world-state storage.",
        "Commander adds the territory module gate, controller masks/badges, contested signaling, standing-supply rows for both factions, controlled-zone counts, authoritative hero capacity, and over-cap text.",
        "Declare disappears when fielded heroes meet capacity; Upgrade and Revive remain available.",
        "The local capture strip shows zone, owner, attacker, progress, remaining seconds, and Awaiting server without predicting a flip.",
        "Capture state clears at INIT context replacement, ordered NEW_WORLD, logout/disconnect, and territory module loss.",
        "The client build passes and commander-map-clinical-check passes 152 assertions.",
    ])

    doc.add_heading("8. Wire connection map", level=1)
    add_table(doc,
        ["Channel", "Carries", "Why it exists"],
        [
            ["837 SMSG_SUI_ZONE_INTEL", "Zone census, controller 0/1/2, contested 0x80, own group units", "Global strategic map; configured controlled zones remain present at zero population"],
            ["839 SMSG_SUI_RTS_STATE", "Mode, module bits, faction Honor/supply/cap, heroes, dungeons", "Campaign snapshot and module-gated UI"],
            ["840/841 RTS action/result", "Declare, Upgrade, Revive and authoritative result", "Intent/result boundary; client never mutates the roster directly"],
            ["842/843 force roster", "Paged faction AiBot force and control eligibility", "Commander selection without converting the force into a WoW party"],
            ["Stock INIT/UPDATE_WORLD_STATE", "Packed local territory capture field 0x53550001", "In-zone capture HUD without allocating another custom opcode"],
        ], [2400, 3880, 3080], font_size=8.8)
    add_table(doc,
        ["Packed bits", "Meaning", "Validation"],
        [
            ["0-1", "Incumbent owner 0/1/2", "No value 3"],
            ["2-3", "Attacker 0/1/2", "Contested requires Alliance or Horde"],
            ["4-5", "Hidden / stable / contested / cooldown", "Phase-specific invariants"],
            ["6-15", "Elapsed progress, 0-1000 permille", "No value above 1000"],
            ["16-31", "Authoritative seconds remaining", "Zero may mean Awaiting server at the client"],
        ], [1600, 4000, 3760], font_size=9)

    doc.add_heading("9. Code impact: where each rule lands", level=1)
    add_table(doc,
        ["Repository / seam", "Required responsibility"],
        [
            ["MangosSuperUI / WorldConfigurationModels", "Sole rts-r3-v1 shape, scalar bounds, keyed zone/hub rules, immutable asset version"],
            ["MangosSuperUI / RtsTerritoryCatalog + WorldStore", "Exact pilot IDs/rows/signatures; collision-safe Character/World artifact DML"],
            ["MangosSuperUI / RtsWorldCreationService", "R3 genesis rows and automatic authored content from stock"],
            ["MangosSuperUI / WorldStateService", "Two-phase resume with semantic preflight and preserved runtime state"],
            ["Core / SuiTerritory", "Validation, queue, state machine, synchronous commit, projections, events, graveyards, local world state"],
            ["Core / SuiRts + SuiHero", "Module latch, tick/lifecycle facade, effective cap, real faction fields, Declare/revive integration"],
            ["Core / GameObject.cpp", "One gated GOOBER banner-use queue hook"],
            ["Core / Player.cpp + AiBot bridge", "Initial/update capture pair and same-map controlled-zone graveyard preference"],
            ["Client / RtsWire.cs", "Typed, atomic, fail-closed zone and packed capture parsing"],
            ["Client / CommanderMap + UI law", "Territory map, standing supply, counts, capacity, over-cap behavior"],
            ["Client / Quest + RtsTerritory HUD + Net", "Raw-first world-state routing, semantic projection, lifecycle fences, capture strip"],
        ], [3260, 6100], font_size=8.7)

    doc.add_heading("10. Concurrency, persistence, and inertness laws", level=1)
    add_bullets(doc, [
        "Map-thread hooks enqueue structural actions; only the main world tick mutates contests, controllers, guards, buffs, or hero roster state.",
        "Territory locks are released before synchronous CharacterDatabase commits. Territory and hero locks are never nested.",
        "There is no territory shutdown flush: every successful flip is already committed before publication.",
        "Core performs runtime DML only. It never creates tables or authors WorldDatabase content.",
        "When RTS mode is false or Territory validation fails, GOOBER use, repop, bots, heroes, wire fields, world states, events, and client UI follow their exact prior behavior.",
        "Client parsing is transactional: a malformed zone or unit block leaves the last complete census/controller/unit snapshot untouched.",
    ])
    add_callout(doc, "OWNER-ONLY BOUNDARY", "No agent action installs or deploys server artifacts, controls a service/process/session, or creates, restores, swaps, or writes a live database/worldstate. Build success is the furthest automated server action.", tone="red")

    doc.add_heading("11. A complete player-facing R3 cycle", level=1)
    add_steps(doc, [
        "The Alliance commander opens Commander and sees Westfall held by Horde, a contested marker absent, standing supply, and current hero capacity.",
        "The commander takes control of an eligible faction AiBot or directs forces through Tier-1 controls and travels to Sentinel Hill.",
        "A faction member uses the neutral banner. The local strip announces the Horde incumbent, Alliance attacker, progress, and remaining time.",
        "If Horde defends through the configured interaction rule, the contest cancels. Otherwise the timer reaches its commit boundary.",
        "Core writes Alliance into superui_zone_control. Only after success do guards swap, supplies and zone counts change, hero capacity update, and Westfall turn blue on Commander.",
        "A newly dead Alliance unit now prefers a Westfall controlled graveyard. Horde uses the vanilla fallback or another friendly controlled zone on that map.",
        "If Alliance later loses Westfall and falls below its fielded hero count, existing heroes remain. Declare is blocked until capacity recovers; Upgrade and Revive continue.",
    ])

    doc.add_heading("12. R3 completion gates and honest status", level=1)
    add_table(doc,
        ["Gate", "Status on 2026-08-16", "Exit proof"],
        [
            ["1A Contract", "Client contract implemented; server/web golden parity still required", "Same constants, mappings, IDs, fixtures across all three repos"],
            ["1B Live discovery", "Blocked on Nico-operated discovery", "Exact banner positions, guard GUID/rows/clones, graveyards, appearance, collision audit"],
            ["2 Web construction", "Not implemented", "Stock input produces exact R3 parked artifacts; collisions fail safely"],
            ["3 Core module", "Not implemented on authoritative Linux checkout", "Release/scripts build plus inertness and malformed-contract checks"],
            ["4 Client", "Implemented and build/clinical verified in this workspace", "Build succeeds; 152 commander assertions pass"],
            ["5 Cross-repo", "Pending Gates 1B-3", "Contract fingerprints and all clinical/source checks agree"],
            ["6 Live acceptance", "Nico-only and pending", "Recorded acceptance matrix passes without MMO regression"],
        ], [1600, 3460, 4300], font_size=8.6)

    doc.add_heading("13. Nico-operated discovery and acceptance", level=1)
    doc.add_heading("13.1 Discovery required before enabling R3", level=2)
    add_bullets(doc, [
        "Physically place and record the final neutral banner pose at all three hubs.",
        "Discover exact native guard creature GUIDs by hub radius and capture their complete source-row signatures.",
        "Define opposing-faction clone rows at matching positions.",
        "Confirm at least one same-map graveyard candidate for each configured zone.",
        "Confirm reserved banner GUIDs 9100108, 9100380, 9100272; guard GUID range 9200000-9299999; template 900001; and event IDs 900-905 are free.",
        "Visually confirm display 6271 and the neutral faction render acceptably.",
    ])
    doc.add_heading("13.2 Live acceptance after Nico deploys", level=2)
    add_bullets(doc, [
        "Stock MMO boot plus normal GOOBER and graveyard behavior.",
        "R3 creation from stock with no manual SQL and correct initial guards/controllers.",
        "Neutral claim, competing assault, defense cancellation, flip, and cooldown at all hubs.",
        "Forced database failure proves no premature publication.",
        "Zero-population held-zone rendering; supplies, count, and cap change together.",
        "Loss below cap preserves heroes and blocks only Declare.",
        "Player, ordinary AiBot, and paid hero graveyard paths.",
        "Restart during a contest restores the old controller; restart after a flip restores the new controller and repairs events.",
        "Both continents, logout/login, disconnect, map transfer, and later MMO load.",
    ])

    doc.add_heading("14. How R4 connects after R3", level=1)
    add_body(doc, "R4 reuses R3's strongest patterns: strict independent module validation, main-thread structural mutation, persist-before-publish control, login-time projection, existing RTS-state rows, and a Nico-operated acceptance boundary.")
    add_table(doc,
        ["R4 mechanic", "Connection to the existing world loop", "Code consequence"],
        [
            ["Entry gate", "Only the non-controlling faction may enter; one live run per faction", "SuiDungeon::CanEnter at area-trigger resolution plus bot travel bridge"],
            ["Clear detection", "Final configured boss kill flips the objective", "Existing Unit::Kill seam enqueues; main tick persists and publishes"],
            ["Faction buff", "Dungeon ownership strengthens the whole army", "Main-thread remove/apply on flip; shared OnPlayerWorldEnter applies on login"],
            ["Boss loot", "Clears accelerate army gearing", "Gated multi-roll at Creature::GenerateLootForBody; default 10, hard cap 16, FFA/no raid-roll fan-out"],
            ["Commander objective", "Strategic map shows controller and live-run flags", "Existing dungeon packet block plus an OBJECTIVES rail"],
        ], [1900, 3800, 3660], font_size=8.7)
    add_callout(doc, "R4 DESIGN FREEZE STILL NEEDED",
        "Choose exact per-dungeon buff spells; define a control flip while the old controller remains inside; freeze run unload/reconnect/timeout laws; and decide whether ownership adds guards or visibility beyond buff and lockout.", tone="gold")

    doc.add_heading("Appendix A. Frozen R3 constants", level=1)
    add_table(doc,
        ["Contract", "Value"],
        [
            ["Territory module bit", "0x04"],
            ["Controller encoding", "0 neutral; 1 Alliance; 2 Horde; OR 0x80 contested"],
            ["Zone-intel row", "Stride 9; legacy stride 8 decodes neutral"],
            ["Faction row", "Stride 26; R3 fills Ore, Skins, Herbs, ControlledZones, HeroSlotCap"],
            ["Capture field", "0x53550001"],
            ["Golden packed capture", "0x002A5DE9 = Alliance incumbent, Horde attacker, contested, 375/1000, 42s"],
            ["R3 scalar keys", "territory.enabled; territory.zones_per_hero_slot; territory.flip_cooldown_ms"],
            ["Pilot profile", "rts-r3-v1 / asset catalog rts-r3-territory-v1"],
        ], [2900, 6460], font_size=9)

    doc.add_heading("Appendix B. Sources and implementation evidence", level=1)
    add_bullets(doc, [
        "docs/systems/SYSTEM_RTS_R3.md - authoritative detailed R3 plan and contract.",
        "RTS_WORLDSTATE_PLAN.md - phased R1-R5 architecture and R4 sketch.",
        "docs/systems/SYSTEM_RTS_R2.md - R2 implementation authority and verification record.",
        "docs/systems/SYSTEM_DATABASE_OVERLAY.md - database ownership and nine-table overlay.",
        "MSUIClient/Net/RtsWire.cs - typed R3 zone/capture parser added in this work.",
        "MSUIClient/GameLoop/Hud/GameLoop.RtsTerritory.cs - local capture lifecycle and HUD added in this work.",
        "MSUIClient/GameLoop/Scene/GameLoop.CommanderMap.cs - territory map, supply, capacity, and over-cap UI added in this work.",
        "tools/commander-map-clinical-check - 152-assertion passing client contract suite.",
    ])

    doc.core_properties.title = "RTS World Architecture and Effective Game Loop"
    doc.core_properties.subject = "Tier-2 RTS R1-R4 system guide and R3 implementation handoff"
    doc.core_properties.author = "Codex for Nico"
    doc.core_properties.keywords = "RTS, VMaNGOS, MSUIClient, MangosSuperUI, R3, territory"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
